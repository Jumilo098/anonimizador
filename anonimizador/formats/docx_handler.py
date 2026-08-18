"""DOCX: se inspecciona MUCHO mas alla del texto visible.

Se abre el ZIP/XML directamente porque python-docx no ve comentarios,
control de cambios, propiedades personalizadas ni relaciones externas.
La salida se RECONSTRUYE desde cero: no se hereda ni una sola propiedad.
"""
from __future__ import annotations

import re
import shutil
import zipfile
from datetime import datetime
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

from ..models import Alerta, Capa, DocumentoExtraido
from .base import ManejadorFormato, hallazgos_de_metadatos, unidad

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

APP_XML_NEUTRO = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/'
    'extended-properties" xmlns:vt="http://schemas.openxmlformats.org/'
    'officeDocument/2006/docPropsVTypes">'
    "<Application>ANONIMIZADOR</Application><Company></Company>"
    "<Manager></Manager><Template></Template><TotalTime>0</TotalTime>"
    "</Properties>"
)

PARTES_A_ELIMINAR = (
    "docProps/custom.xml",
    "word/comments.xml",
    "word/commentsExtended.xml",
    "word/commentsIds.xml",
    "word/commentsExtensible.xml",
    "word/people.xml",
    "word/vbaProject.bin",
    "word/vbaData.xml",
    "docProps/thumbnail.jpeg",
    "docProps/thumbnail.emf",
    "docProps/thumbnail.wmf",
    "customXml/item1.xml",
    "customXml/itemProps1.xml",
    "customXml/_rels/item1.xml.rels",
)


def _sin_ns(tag: str) -> str:
    return tag.split("}")[-1]


def _texto_elemento(el, incluir_borrados=False) -> str:
    """Texto de un w:p incluyendo inserciones de control de cambios."""
    partes = []
    for nodo in el.iter():
        t = _sin_ns(nodo.tag)
        if t == "t" and nodo.text:
            partes.append(nodo.text)
        elif t == "tab":
            partes.append("\t")
        elif t in ("br", "cr"):
            partes.append("\n")
        elif incluir_borrados and t == "delText" and nodo.text:
            partes.append(nodo.text)
    return "".join(partes)


def _texto_borrado(el) -> str:
    partes = [n.text for n in el.iter() if _sin_ns(n.tag) == "delText" and n.text]
    return "".join(partes)


class ManejadorDocx(ManejadorFormato):
    extensiones = (".docx",)
    nombre = "docx"
    reconstruccion = "completa"
    capas_no_copiadas = frozenset({Capa.COMENTARIO, Capa.CONTROL_CAMBIOS})

    # -- extraccion --------------------------------------------------------
    def extraer(self, ruta: Path) -> DocumentoExtraido:
        ruta = Path(ruta)
        doc = DocumentoExtraido(formato="docx")
        try:
            documento = Document(str(ruta))
        except Exception as exc:  # archivo corrupto: se reporta, no se ejecuta
            doc.alertas.append(
                Alerta("critica", "DOCX_ILEGIBLE",
                       "No se pudo abrir el DOCX", str(exc)[:200])
            )
            return doc

        estructura = []
        idx = 0
        cuerpo = documento.element.body
        for hijo in cuerpo.iterchildren():
            etiqueta = _sin_ns(hijo.tag)
            if etiqueta == "p":
                idx += 1
                uid = "p%d" % idx
                texto = _texto_elemento(hijo)
                estilo = ""
                ppr = hijo.find(W + "pPr")
                if ppr is not None:
                    pstyle = ppr.find(W + "pStyle")
                    if pstyle is not None:
                        estilo = pstyle.get(W + "val") or ""
                doc.unidades.append(
                    unidad(uid, texto, Capa.CONTENIDO, "parrafo[%d]" % idx,
                           estilo=estilo)
                )
                estructura.append({"tipo": "parrafo", "uid": uid, "estilo": estilo})
                borrado = _texto_borrado(hijo)
                if borrado.strip():
                    doc.unidades.append(
                        unidad(uid + "_del", borrado, Capa.CONTROL_CAMBIOS,
                               "control_de_cambios[%d]" % idx, editable=False)
                    )
            elif etiqueta == "tbl":
                idx += 1
                filas = []
                for f, fila in enumerate(hijo.findall(W + "tr")):
                    celdas = []
                    textos = [_texto_elemento(c) for c in fila.findall(W + "tc")]
                    for c, celda in enumerate(fila.findall(W + "tc")):
                        uid = "t%d_%d_%d" % (idx, f, c)
                        # En las fichas clinicas el par etiqueta/valor vive en
                        # celdas distintas ("Municipio" | "Villa Robleda"), asi
                        # que la etiqueta de la celda anterior viaja como
                        # contexto para que el detector pueda usarla.
                        doc.unidades.append(
                            unidad(uid, textos[c], Capa.TABLA,
                                   "tabla[%d].celda[%d,%d]" % (idx, f, c),
                                   etiqueta_previa=textos[c - 1] if c > 0 else "")
                        )
                        celdas.append(uid)
                    filas.append(celdas)
                estructura.append({"tipo": "tabla", "filas": filas})

        # encabezados y pies (todas las secciones, todos los tipos)
        for s, seccion in enumerate(documento.sections):
            for atributo, etiqueta, capa in (
                ("header", "encabezado", Capa.ENCABEZADO),
                ("first_page_header", "encabezado_primera", Capa.ENCABEZADO),
                ("even_page_header", "encabezado_par", Capa.ENCABEZADO),
                ("footer", "pie", Capa.PIE),
                ("first_page_footer", "pie_primera", Capa.PIE),
                ("even_page_footer", "pie_par", Capa.PIE),
            ):
                try:
                    contenedor = getattr(seccion, atributo, None)
                    texto = "\n".join(
                        p.text for p in contenedor.paragraphs
                    ).strip() if contenedor is not None else ""
                except Exception:
                    texto = ""
                if texto:
                    uid = "%s_%d" % (etiqueta, s)
                    doc.unidades.append(
                        unidad(uid, texto, capa, "%s[seccion %d]" % (etiqueta, s))
                    )

        # partes internas del ZIP
        doc.metadatos, extra, alertas = self._inspeccionar_zip(ruta, doc)
        doc.alertas.extend(alertas)
        doc.unidades.extend(extra)
        doc.unidades.append(
            unidad("nombre_archivo", ruta.stem, Capa.NOMBRE_ARCHIVO,
                   "nombre del archivo")
        )
        doc.hallazgos_tecnicos = hallazgos_de_metadatos(doc.metadatos)
        doc.info["estructura"] = estructura
        doc.info["parrafos"] = sum(1 for e in estructura if e["tipo"] == "parrafo")
        doc.info["tablas"] = sum(1 for e in estructura if e["tipo"] == "tabla")
        return doc

    def _inspeccionar_zip(self, ruta: Path, doc: DocumentoExtraido):
        metadatos = {}
        unidades = []
        alertas = []
        activos = []
        with zipfile.ZipFile(ruta) as z:
            nombres = z.namelist()
            for parte, prefijo in (("docProps/core.xml", "core"),
                                   ("docProps/app.xml", "app")):
                if parte in nombres:
                    try:
                        raiz = ET.fromstring(z.read(parte))
                        for nodo in raiz:
                            valor = (nodo.text or "").strip()
                            if valor:
                                metadatos[prefijo + ":" + _sin_ns(nodo.tag)] = valor
                    except ET.ParseError:
                        alertas.append(Alerta("advertencia", "XML_INVALIDO",
                                              "No se pudo leer " + parte))
            if "docProps/custom.xml" in nombres:
                try:
                    raiz = ET.fromstring(z.read("docProps/custom.xml"))
                    for prop in raiz:
                        nombre = prop.get("name") or "propiedad"
                        valor = "".join(x.text or "" for x in prop).strip()
                        if valor:
                            metadatos["custom:" + nombre] = valor
                except ET.ParseError:
                    pass
            # comentarios
            if "word/comments.xml" in nombres:
                try:
                    raiz = ET.fromstring(z.read("word/comments.xml"))
                    for i, com in enumerate(raiz):
                        autor = com.get(W + "author") or ""
                        iniciales = com.get(W + "initials") or ""
                        texto = _texto_elemento(com)
                        if autor:
                            metadatos["comentario_autor[%d]" % i] = autor
                        if iniciales:
                            metadatos["comentario_iniciales[%d]" % i] = iniciales
                        if texto.strip():
                            unidades.append(
                                unidad("comentario_%d" % i, texto, Capa.COMENTARIO,
                                       "word/comments.xml[%d]" % i, editable=False)
                            )
                except ET.ParseError:
                    pass
            # notas al pie
            for parte, etiqueta in (("word/footnotes.xml", "nota_pie"),
                                    ("word/endnotes.xml", "nota_final")):
                if parte in nombres:
                    try:
                        raiz = ET.fromstring(z.read(parte))
                        for i, nota in enumerate(raiz):
                            texto = _texto_elemento(nota).strip()
                            if len(texto) > 1:
                                unidades.append(
                                    unidad("%s_%d" % (etiqueta, i), texto,
                                           Capa.CONTENIDO, parte + "[%d]" % i,
                                           editable=False)
                                )
                    except ET.ParseError:
                        pass
            # autores de control de cambios
            try:
                cuerpo = z.read("word/document.xml").decode("utf-8", "replace")
                raiz = ET.fromstring(cuerpo)
                autores = set()
                for nodo in raiz.iter():
                    if _sin_ns(nodo.tag) in ("ins", "del", "moveFrom", "moveTo"):
                        a = nodo.get(W + "author")
                        if a:
                            autores.add(a)
                for i, a in enumerate(sorted(autores)):
                    metadatos["control_cambios_autor[%d]" % i] = a
                if autores:
                    alertas.append(
                        Alerta("advertencia", "CONTROL_CAMBIOS",
                               "El documento traia control de cambios con autores.",
                               "Autores detectados: %d" % len(autores))
                    )
            except (KeyError, ET.ParseError):
                pass
            # relaciones externas
            for parte in [n for n in nombres if n.endswith(".rels")]:
                try:
                    raiz = ET.fromstring(z.read(parte))
                    for rel in raiz:
                        if rel.get("TargetMode") == "External":
                            metadatos["rel_externa:" + parte] = rel.get("Target", "")
                except ET.ParseError:
                    pass
            # objetos y medios
            medios = [n for n in nombres if n.startswith("word/media/")]
            embebidos = [n for n in nombres if n.startswith("word/embeddings/")]
            macros = [n for n in nombres if "vbaProject" in n]
            for n in medios:
                activos.append({"tipo": "imagen", "ruta": n,
                                "bytes": z.getinfo(n).file_size})
            for n in embebidos:
                activos.append({"tipo": "objeto_embebido", "ruta": n,
                                "bytes": z.getinfo(n).file_size})
            if medios:
                alertas.append(
                    Alerta("advertencia", "IMAGENES_EN_DOCX",
                           "El DOCX contiene %d imagen(es)." % len(medios),
                           "En V1 NO se transfieren a la version reconstruida "
                           "y no se analizan sus pixeles. Revise cada imagen aparte.")
                )
            if embebidos:
                alertas.append(
                    Alerta("critica", "OBJETOS_EMBEBIDOS",
                           "El DOCX contiene %d objeto(s) embebido(s)." % len(embebidos),
                           "No se abren ni se ejecutan. No se transfieren al resultado.")
                )
            if macros:
                alertas.append(
                    Alerta("critica", "MACROS",
                           "El archivo contiene proyecto de macros (VBA).",
                           "NO se ejecuta. La reconstruccion lo elimina.")
                )
        doc.activos = activos
        return metadatos, unidades, alertas

    # -- reconstruccion ----------------------------------------------------
    def reconstruir(self, extraido, unidades_nuevas, destino: Path, contexto=None):
        mapa = {u.uid: u for u in unidades_nuevas}
        nuevo = Document()
        partes_texto = []

        def parrafo(texto, estilo=""):
            estilo_norm = (estilo or "").lower()
            if estilo_norm.startswith("heading") or estilo_norm.startswith("titulo") \
                    or estilo_norm.startswith("ttulo") or estilo_norm.startswith("title"):
                try:
                    nivel = int("".join(c for c in estilo_norm if c.isdigit()) or 1)
                except ValueError:
                    nivel = 1
                try:
                    nuevo.add_heading(texto, min(max(nivel, 0), 4))
                    return
                except Exception:
                    p = nuevo.add_paragraph()
                    p.add_run(texto).bold = True
                    return
            nuevo.add_paragraph(texto)

        for bloque in extraido.info.get("estructura", []):
            if bloque["tipo"] == "parrafo":
                u = mapa.get(bloque["uid"])
                texto = u.texto if u else ""
                parrafo(texto, bloque.get("estilo", ""))
                partes_texto.append(texto)
            else:
                filas = bloque["filas"]
                if not filas:
                    continue
                tabla = nuevo.add_table(rows=len(filas), cols=max(len(f) for f in filas))
                try:
                    tabla.style = "Table Grid"
                except Exception:
                    pass
                for i, fila in enumerate(filas):
                    for j, uid in enumerate(fila):
                        u = mapa.get(uid)
                        texto = u.texto if u else ""
                        tabla.cell(i, j).text = texto
                        partes_texto.append(texto)

        descartadas = []
        anexos = [
            (Capa.ENCABEZADO, "ANEXO A - ENCABEZADOS (extraidos y minimizados)"),
            (Capa.PIE, "ANEXO B - PIES DE PAGINA (extraidos y minimizados)"),
        ]
        for capa, titulo in anexos:
            textos = [u.texto.strip() for u in unidades_nuevas
                      if u.capa == capa and u.texto.strip()]
            if not textos:
                continue
            nuevo.add_paragraph("")
            parrafo(titulo, "heading2")
            partes_texto.append(titulo)
            for t in textos:
                nuevo.add_paragraph(t)
                partes_texto.append(t)

        for capa, etiqueta in ((Capa.COMENTARIO, "comentarios"),
                               (Capa.CONTROL_CAMBIOS, "control de cambios")):
            if any(u.capa == capa and u.texto.strip() for u in unidades_nuevas):
                descartadas.append(etiqueta)

        # propiedades neutras
        cp = nuevo.core_properties
        cp.author = ""
        cp.last_modified_by = ""
        cp.title = "Documento desidentificado"
        cp.subject = ""
        cp.keywords = ""
        cp.comments = "Generado por ANONIMIZADOR. Requiere revision humana."
        cp.category = ""
        cp.content_status = "DESIDENTIFICADO - PENDIENTE DE REVISION"
        cp.identifier = ""
        cp.language = ""
        cp.version = ""
        cp.revision = 1
        neutro = datetime(2000, 1, 1, 0, 0, 0)
        cp.created = neutro
        cp.modified = neutro
        cp.last_printed = neutro

        destino = Path(destino)
        nuevo.save(str(destino))
        self._sanear_zip(destino)

        return {
            "texto_esperado": "\n".join(partes_texto),
            "capas_descartadas": descartadas,
            "notas": [
                "Documento reconstruido desde cero con python-docx.",
                "No se hereda ninguna propiedad del original.",
                "Comentarios, control de cambios, macros, objetos embebidos e "
                "imagenes NO se transfieren.",
            ],
        }

    def _sanear_zip(self, ruta: Path):
        """Reescribe el ZIP quitando partes que no deben existir jamas."""
        tmp = ruta.with_suffix(".tmp.docx")
        with zipfile.ZipFile(ruta) as origen:
            entradas = [i for i in origen.infolist()]
            datos = {i.filename: origen.read(i.filename) for i in entradas}
        eliminadas = [p for p in PARTES_A_ELIMINAR if datos.pop(p, None) is not None]
        if eliminadas and "[Content_Types].xml" in datos:
            ct = datos["[Content_Types].xml"].decode("utf-8", "replace")
            for parte in eliminadas:
                ct = re.sub(
                    r"<Override[^>]*PartName=\"/" + re.escape(parte) + r"\"[^>]*/>",
                    "", ct,
                )
            datos["[Content_Types].xml"] = ct.encode("utf-8")
        if eliminadas:
            # tambien se quitan las relaciones que apuntaban a esas partes
            objetivos = {p.split("/")[-1] for p in eliminadas}
            for nombre in [n for n in datos if n.endswith(".rels")]:
                xml = datos[nombre].decode("utf-8", "replace")
                for objetivo in objetivos:
                    xml = re.sub(
                        r"<Relationship[^>]*Target=\"[^\"]*"
                        + re.escape(objetivo) + r"\"[^>]*/>",
                        "", xml,
                    )
                datos[nombre] = xml.encode("utf-8")
        if "docProps/app.xml" in datos:
            datos["docProps/app.xml"] = APP_XML_NEUTRO.encode("utf-8")
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as salida:
            for nombre, contenido in datos.items():
                salida.writestr(nombre, contenido)
        shutil.move(str(tmp), str(ruta))
