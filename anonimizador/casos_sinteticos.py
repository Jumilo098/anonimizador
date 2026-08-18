"""Generador de casos clinicos COMPLETAMENTE FICTICIOS.

Ni un solo dato de estos archivos corresponde a una persona real.
Sirven para la demostracion en clase y para los tests automaticos.
"""
from __future__ import annotations

import zipfile
from pathlib import Path

from .config import SAMPLES_DIR

BANNER = "CASO SINTETICO - PACIENTE FICTICIO - NO CORRESPONDE A NINGUNA PERSONA REAL"

# ---------------------------------------------------------------------------
# Texto base compartido (identificadores + datos clinicos que deben sobrevivir)
# ---------------------------------------------------------------------------
NOTA_EVOLUCION = """{banner}

HOSPITAL UNIVERSITARIO SAN RAFAEL - Servicio de Gastroenterologia
NIT 900.123.456-7   Sede Bogota

NOTA DE EVOLUCION

Paciente: Maria Fernanda Lopez Quintero
Documento: CC 52.123.456
Historia clinica: HC 0099887
Fecha de nacimiento: 14/02/1960
Edad: 64 anos
Sexo: femenino
Ocupacion: docente
Direccion: Calle 127 # 15-45, Bogota
Telefono: 320 555 1234
Correo: mafe.lopez1960@correo-ficticio.co
EPS: Sanitas   Poliza: POL-778812
Acompanante: Jorge Enrique Lopez (hermano)
Fecha de consulta: 12/03/2024

MOTIVO DE CONSULTA
Astenia de 3 meses de evolucion y melenas intermitentes.

ANTECEDENTES
Hipertension arterial diagnosticada hace 8 anos, en manejo con losartan 50 mg
cada 12 horas. No refiere alergias. No consumo de alcohol.

EXAMEN FISICO
TA 118/76 mmHg, FC 88 lpm, FR 18 rpm, temperatura 36.8 C, SatO2 95%.
Peso 61.5 kg, talla 158 cm, IMC 24.6 kg/m2.
Palidez mucocutanea. Abdomen blando, sin masas palpables.

LABORATORIOS del 10/03/2024
Hemoglobina: 8.9 g/dL
Hematocrito: 28.4 %
VCM: 71 fL
Ferritina: 6 ng/mL
Creatinina: 0.9 mg/dL
Leucocitos: 7.200 /uL
Plaquetas: 310000 /uL

COLONOSCOPIA del 11/03/2024
Se identifica lesion ulcerada de 3.2 x 2.1 cm en colon ascendente, a 65 cm del
margen anal. Se toman biopsias. No se observan lesiones en colon sigmoide ni en
recto.

IMPRESION DIAGNOSTICA
1. Anemia ferropenica microcitica hipocromica secundaria a sangrado digestivo.
2. Lesion en colon ascendente, probable adenocarcinoma. Pendiente patologia.
3. Hipertension arterial controlada.

PLAN
- Sulfato ferroso 300 mg via oral cada 12 horas.
- TAC de abdomen contrastado el 15/03/2024.
- Control con resultado de biopsia el 20/03/2024.
- Se explica el caso a la paciente y al acompanante.

Elaborado por: Dr. Fernando Gil Bermudez - Registro medico 12345
Revisado por: Dra. Claudia Restrepo Arias
""".format(banner=BANNER)


def _escribir_txt(destino: Path) -> Path:
    destino.write_text(NOTA_EVOLUCION, encoding="utf-8")
    return destino


# ---------------------------------------------------------------------------
# DOCX con TODAS las capas sucias
# ---------------------------------------------------------------------------
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

CORE_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
<dc:title>Historia clinica Maria Fernanda Lopez Quintero</dc:title>
<dc:subject>HC 0099887 - Gastroenterologia</dc:subject>
<dc:creator>Dr. Fernando Gil Bermudez</dc:creator>
<cp:keywords>lopez quintero; cc 52123456; sanitas</cp:keywords>
<dc:description>Revisar telefono 320 555 1234</dc:description>
<cp:lastModifiedBy>claudia.restrepo@hospital-ficticio.co</cp:lastModifiedBy>
<cp:revision>7</cp:revision>
<dcterms:created xsi:type="dcterms:W3CDTF">2024-03-12T14:03:00Z</dcterms:created>
<dcterms:modified xsi:type="dcterms:W3CDTF">2024-03-13T09:41:00Z</dcterms:modified>
<cp:category>Historias clinicas</cp:category>
</cp:coreProperties>"""

APP_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
<Application>Microsoft Office Word</Application>
<Company>Hospital Universitario San Rafael</Company>
<Manager>Dra. Claudia Restrepo Arias</Manager>
<Template>PlantillaHC_SanRafael.dotx</Template>
<TotalTime>128</TotalTime>
</Properties>"""

CUSTOM_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="2" name="NumeroPaciente"><vt:lpwstr>PAC-0099887</vt:lpwstr></property>
<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="3" name="MedicoTratante"><vt:lpwstr>Dr. Fernando Gil Bermudez</vt:lpwstr></property>
<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="4" name="CorreoContacto"><vt:lpwstr>mafe.lopez1960@correo-ficticio.co</vt:lpwstr></property>
</Properties>"""

COMMENTS_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="%s">
<w:comment w:id="1" w:author="Dra. Claudia Restrepo Arias" w:date="2024-03-13T09:20:00Z" w:initials="CRA">
<w:p><w:r><w:t>Confirmar telefono del acudiente Jorge Enrique Lopez: 315 444 9988</w:t></w:r></w:p>
</w:comment>
</w:comments>""" % W_NS

PARRAFO_SUCIO = (
    '<w:p>'
    '<w:commentRangeStart w:id="1"/>'
    '<w:ins w:id="101" w:author="Dr. Fernando Gil Bermudez" w:date="2024-03-13T09:00:00Z">'
    '<w:r><w:t>Se agrega: la paciente refiere astenia progresiva desde enero.</w:t></w:r>'
    '</w:ins>'
    '<w:del w:id="102" w:author="Dra. Claudia Restrepo Arias" w:date="2024-03-13T09:05:00Z">'
    '<w:r><w:delText>Texto borrado: paciente Maria Fernanda Lopez Quintero, '
    'CC 52.123.456, telefono 320 555 1234.</w:delText></w:r>'
    '</w:del>'
    '<w:commentRangeEnd w:id="1"/>'
    '<w:r><w:commentReference w:id="1"/></w:r>'
    '</w:p>'
)

REL_CUSTOM = ('<Relationship Id="rIdCustomProps" Type="http://schemas.openxmlformats.org/'
              'officeDocument/2006/relationships/custom-properties" '
              'Target="docProps/custom.xml"/>')
REL_COMMENTS = ('<Relationship Id="rIdComments" Type="http://schemas.openxmlformats.org/'
                'officeDocument/2006/relationships/comments" Target="comments.xml"/>')
CT_CUSTOM = ('<Override PartName="/docProps/custom.xml" ContentType="application/'
             'vnd.openxmlformats-officedocument.custom-properties+xml"/>')
CT_COMMENTS = ('<Override PartName="/word/comments.xml" ContentType="application/'
               'vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>')


def _escribir_docx(destino: Path) -> Path:
    from docx import Document

    doc = Document()
    seccion = doc.sections[0]
    seccion.header.paragraphs[0].text = (
        "Hospital Universitario San Rafael - HC 0099887 - Maria Fernanda Lopez"
    )
    seccion.footer.paragraphs[0].text = (
        "Impreso por claudia.restrepo@hospital-ficticio.co - 13/03/2024"
    )

    doc.add_heading("Nota de evolucion", 1)
    for bloque in NOTA_EVOLUCION.split("\n\n"):
        bloque = bloque.strip()
        if not bloque:
            continue
        if bloque.isupper() and len(bloque) < 60:
            doc.add_heading(bloque.title(), 2)
        else:
            doc.add_paragraph(bloque)

    doc.add_heading("Laboratorios (tabla)", 2)
    tabla = doc.add_table(rows=5, cols=3)
    datos = [
        ["Analito", "Resultado", "Referencia"],
        ["Hemoglobina", "8.9 g/dL", "12.0 - 16.0 g/dL"],
        ["Ferritina", "6 ng/mL", "15 - 150 ng/mL"],
        ["Creatinina", "0.9 mg/dL", "0.6 - 1.1 mg/dL"],
        ["Paciente", "Maria Fernanda Lopez Quintero - CC 52.123.456", ""],
    ]
    for i, fila in enumerate(datos):
        for j, valor in enumerate(fila):
            tabla.cell(i, j).text = valor

    doc.save(str(destino))
    _ensuciar_docx(destino)
    return destino


def _ensuciar_docx(ruta: Path):
    """Inyecta metadatos, propiedades personalizadas, comentarios y control de
    cambios: exactamente lo que un DOCX real arrastra sin que nadie lo vea."""
    with zipfile.ZipFile(ruta) as z:
        datos = {i.filename: z.read(i.filename) for i in z.infolist()}

    datos["docProps/core.xml"] = CORE_XML.encode("utf-8")
    datos["docProps/app.xml"] = APP_XML.encode("utf-8")
    datos["docProps/custom.xml"] = CUSTOM_XML.encode("utf-8")
    datos["word/comments.xml"] = COMMENTS_XML.encode("utf-8")

    ct = datos["[Content_Types].xml"].decode("utf-8")
    if "custom.xml" not in ct:
        ct = ct.replace("</Types>", CT_CUSTOM + CT_COMMENTS + "</Types>")
    datos["[Content_Types].xml"] = ct.encode("utf-8")

    rels = datos["_rels/.rels"].decode("utf-8")
    if "custom.xml" not in rels:
        rels = rels.replace("</Relationships>", REL_CUSTOM + "</Relationships>")
    datos["_rels/.rels"] = rels.encode("utf-8")

    drels = datos["word/_rels/document.xml.rels"].decode("utf-8")
    if "comments.xml" not in drels:
        drels = drels.replace("</Relationships>", REL_COMMENTS + "</Relationships>")
    datos["word/_rels/document.xml.rels"] = drels.encode("utf-8")

    doc = datos["word/document.xml"].decode("utf-8")
    corte = doc.rfind("<w:sectPr")
    if corte == -1:
        corte = doc.rfind("</w:body>")
    doc = doc[:corte] + PARRAFO_SUCIO + doc[corte:]
    datos["word/document.xml"] = doc.encode("utf-8")

    with zipfile.ZipFile(ruta, "w", zipfile.ZIP_DEFLATED) as z:
        for nombre, contenido in datos.items():
            z.writestr(nombre, contenido)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def _escribir_pdf(destino: Path) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfgen import canvas as rl_canvas

    ancho, alto = A4
    c = rl_canvas.Canvas(str(destino), pagesize=A4)
    c.setTitle("Informe de patologia - Maria Fernanda Lopez Quintero")
    c.setAuthor("Dr. Fernando Gil Bermudez")
    c.setSubject("HC 0099887 - CC 52.123.456")
    c.setKeywords("lopez quintero, sanitas, hospital san rafael")
    c.setCreator("Sistema HIS Hospital San Rafael v4.2")

    texto = [
        BANNER,
        "",
        "HOSPITAL UNIVERSITARIO SAN RAFAEL - Laboratorio de Patologia",
        "Calle 127 # 15-45, Bogota - Telefono: 601 555 8899",
        "",
        "INFORME DE PATOLOGIA",
        "",
        "Paciente: Maria Fernanda Lopez Quintero",
        "Documento: CC 52.123.456    Historia clinica: HC 0099887",
        "Edad: 64 anos    Sexo: femenino    Ocupacion: docente",
        "EPS: Sanitas    Orden No. 88123456",
        "Correo: mafe.lopez1960@correo-ficticio.co",
        "Fecha de la muestra: 11/03/2024    Fecha del informe: 18/03/2024",
        "",
        "MUESTRA",
        "Biopsia endoscopica de colon ascendente (4 fragmentos).",
        "",
        "DESCRIPCION MACROSCOPICA",
        "Fragmentos parduzcos que en conjunto miden 0.8 x 0.5 x 0.3 cm.",
        "",
        "DESCRIPCION MICROSCOPICA",
        "Mucosa colonica con proliferacion glandular atipica, nucleos",
        "hipercromaticos y mitosis frecuentes. Se observa invasion de la",
        "submucosa. No se observa invasion linfovascular.",
        "",
        "DIAGNOSTICO",
        "Adenocarcinoma bien diferenciado de colon ascendente.",
        "Margenes de la biopsia no evaluables por el tipo de muestra.",
        "",
        "Patologo: Dra. Claudia Restrepo Arias - Registro medico 45678",
        "Firmado digitalmente por: Claudia Restrepo Arias",
    ]
    y = alto - 2 * cm
    c.setFont("Helvetica", 10)
    for linea in texto:
        c.drawString(2 * cm, y, linea)
        y -= 14
    c.showPage()
    c.save()
    return destino


# ---------------------------------------------------------------------------
# Imagen con texto quemado en los pixeles + QR sintetico
# ---------------------------------------------------------------------------
def _escribir_png(destino: Path) -> Path:
    from PIL import Image, ImageDraw, ImageFont
    from PIL.PngImagePlugin import PngInfo

    ancho, alto = 900, 620
    img = Image.new("RGB", (ancho, alto), (245, 245, 245))
    d = ImageDraw.Draw(img)

    def fuente(tam):
        try:
            return ImageFont.load_default(size=tam)
        except TypeError:
            return ImageFont.load_default()

    d.rectangle((0, 0, ancho, 130), fill=(30, 60, 110))
    d.text((20, 14), "HOSPITAL UNIVERSITARIO SAN RAFAEL", font=fuente(20),
           fill=(255, 255, 255))
    d.text((20, 44), "Paciente: MARIA FERNANDA LOPEZ QUINTERO", font=fuente(18),
           fill=(255, 255, 255))
    d.text((20, 70), "HC 0099887   CC 52.123.456   F. estudio: 11/03/2024",
           font=fuente(16), fill=(255, 255, 255))
    d.text((20, 96), "Estudio: TAC de abdomen contrastado", font=fuente(16),
           fill=(255, 255, 255))

    # "area clinica": una simulacion de imagen medica que NO se debe tocar
    d.rectangle((60, 160, 620, 580), fill=(20, 20, 20))
    for i in range(10):
        d.ellipse((160 + i * 12, 250 + i * 8, 460 - i * 10, 500 - i * 12),
                  outline=(120 + i * 10, 120 + i * 10, 120 + i * 10))
    d.text((70, 168), "SERIE 3 / IMG 42", font=fuente(14), fill=(220, 220, 220))

    try:
        import qrcode
        qr = qrcode.QRCode(box_size=4, border=2)
        qr.add_data("https://his.hospital-ficticio.co/paciente/0099887")
        qr.make(fit=True)
        qr_img = qr.make_image(fill_color="black", back_color="white").convert("RGB")
        img.paste(qr_img, (660, 170))
    except Exception:
        d.rectangle((660, 170, 820, 330), fill=(0, 0, 0))

    d.text((660, 350), "Etiqueta: PAC-0099887", font=fuente(14), fill=(20, 20, 20))
    d.text((660, 375), "Tel: 320 555 1234", font=fuente(14), fill=(20, 20, 20))
    d.text((20, 596), BANNER, font=fuente(13), fill=(120, 120, 120))

    meta = PngInfo()
    meta.add_text("Author", "Dr. Fernando Gil Bermudez")
    meta.add_text("Comment", "Paciente Maria Fernanda Lopez Quintero HC 0099887")
    meta.add_text("Software", "Estacion PACS Hospital San Rafael")
    img.save(destino, "PNG", pnginfo=meta)
    return destino


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------
def _escribir_xlsx(destino: Path) -> Path:
    from openpyxl import Workbook

    wb = Workbook()
    hoja = wb.active
    hoja.title = "HC Lopez Quintero"
    filas = [
        [BANNER, "", "", ""],
        ["Paciente", "Maria Fernanda Lopez Quintero", "", ""],
        ["Documento", "CC 52.123.456", "Historia clinica", "HC 0099887"],
        ["Correo", "mafe.lopez1960@correo-ficticio.co", "Telefono", "320 555 1234"],
        ["Edad", 64, "Ocupacion", "docente"],
        ["Institucion", "Hospital Universitario San Rafael", "EPS", "Sanitas"],
        ["", "", "", ""],
        ["Analito", "Resultado", "Unidad", "Referencia"],
        ["Hemoglobina", 8.9, "g/dL", "12.0 - 16.0"],
        ["Hematocrito", 28.4, "%", "36 - 46"],
        ["Ferritina", 6, "ng/mL", "15 - 150"],
        ["Creatinina", 0.9, "mg/dL", "0.6 - 1.1"],
        ["Plaquetas", 310000, "/uL", "150000 - 450000"],
    ]
    for fila in filas:
        hoja.append(fila)
    props = wb.properties
    props.creator = "Dr. Fernando Gil Bermudez"
    props.lastModifiedBy = "claudia.restrepo@hospital-ficticio.co"
    props.title = "Laboratorios Maria Fernanda Lopez Quintero"
    props.keywords = "HC 0099887; CC 52123456"
    wb.save(str(destino))
    return destino


# ---------------------------------------------------------------------------
CASOS = {
    "caso_01_nota_evolucion.txt": _escribir_txt,
    "caso_02_historia_clinica.docx": _escribir_docx,
    "caso_03_informe_patologia.pdf": _escribir_pdf,
    "caso_04_etiqueta_estudio.png": _escribir_png,
    "caso_05_laboratorios.xlsx": _escribir_xlsx,
}


def generar(directorio=None, forzar=False):
    """Crea todos los casos sinteticos y devuelve sus rutas."""
    base = Path(directorio or SAMPLES_DIR)
    base.mkdir(parents=True, exist_ok=True)
    rutas = []
    for nombre, funcion in CASOS.items():
        destino = base / nombre
        if forzar or not destino.exists():
            funcion(destino)
        rutas.append(destino)
    (base / "LEEME.txt").write_text(
        BANNER + "\n\nTodos los archivos de esta carpeta son ficticios y se generan\n"
        "automaticamente con `python -m anonimizador.casos_sinteticos`.\n"
        "Nunca guarde aqui documentos de pacientes reales.\n",
        encoding="utf-8",
    )
    return rutas


if __name__ == "__main__":
    for r in generar(forzar=True):
        print("generado:", r)
